from io import BytesIO
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt  # для управления отступами между абзацами

from .schemas import ReportRequest, Section


def _add_heading(document: Document, text: str, level: int = 1):
    return document.add_heading(text, level=level)


def _add_paragraph(document: Document, text: str, bold: bool = False, italic: bool = False, align=None, style: str | None = None):
    p = document.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if align is not None:
        p.alignment = align
    return p


def _add_section_as_list(document: Document, section: Section):
    """
    Раздел выводится построчно:
    • строка с формулировкой проблемы (якорь для комментария),
    • сразу следом строка с жирной меткой 'Исправление:' и текстом решения.
    Между этими двумя строками НЕТ дополнительного отступа.
    """
    _add_heading(document, section.part, level=1)

    for inst in section.instances or []:
        # 1) Проблема — буллет, к этому run привязываем комментарий
        p1 = document.add_paragraph(style="List Bullet")
        anchor_run = p1.add_run(inst.what_is_incorrect)

        # убираем отступ после проблемы
        pf1 = p1.paragraph_format
        pf1.space_after = Pt(0)

        # Комментарий (LLM ID + Риски) к якорю
        comment = document.add_comment(
            runs=[anchor_run],
            text="",
            author="Auto",
            initials="AG",
        )
        # Строки внутри комментария
        comment.paragraphs[0].add_run(f"LLM ID: {inst.llm_id}")
        comment.add_paragraph().add_run(f"Риски: {inst.risks}")

        # 2) Исправление — отдельный абзац, без отступов до/после
        p2 = document.add_paragraph()
        pf2 = p2.paragraph_format
        pf2.space_before = Pt(0)
        pf2.space_after = Pt(0)

        label = p2.add_run("Исправление:")
        label.bold = True
        p2.add_run(f" {inst.how_to_fix}")

    # небольшой отступ между СЕКЦИЯМИ (оставляем)
    document.add_paragraph("")


def build_report(req: ReportRequest) -> bytes:
    """Генерирует .docx-отчёт из JSON и возвращает байты (без сводки, приоритетов и примечаний)."""
    doc = Document()

    # Титул
    title = doc.add_heading(f"Отчёт об ошибках в ТЗ: {req.doc_title}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_paragraph(doc, f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M')}", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Секции
    for section in req.sections or []:
        _add_section_as_list(doc, section)

    # Сохранение
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
